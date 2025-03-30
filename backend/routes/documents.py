import os
import logging
import traceback
import tempfile
import base64
from io import BytesIO
import PyPDF2
from flask import request, jsonify
import googleapiclient.discovery
from googleapiclient.http import MediaIoBaseDownload
from utils import create_credentials_from_token, sanitize_html

logger = logging.getLogger(__name__)

def register_routes(app):
    @app.route('/api/extract-text', methods=['POST'])
    def extract_text():
        """简化的文本提取端点 - 更可靠方式提取各种文档的文本"""
        try:
            data = request.json
            if not data:
                return jsonify({"error": "缺少请求数据"}), 400
                
            file_id = data.get('file_id')
            if not file_id:
                return jsonify({"error": "缺少file_id参数"}), 400
                
            # 创建凭证
            credentials = create_credentials_from_token(data)
            drive_service = googleapiclient.discovery.build('drive', 'v3', credentials=credentials)
            
            # 获取文件元数据
            file_metadata = drive_service.files().get(fileId=file_id, fields="name,mimeType").execute()
            mime_type = file_metadata.get('mimeType', '')
            file_name = file_metadata.get('name', 'unknown')
            
            extracted_text = ""
            
            # 根据不同的文件类型提取文本
            if mime_type == 'application/vnd.google-apps.document':
                # Google文档 - 导出为纯文本
                response = drive_service.files().export(
                    fileId=file_id,
                    mimeType='text/plain'
                ).execute()
                
                if isinstance(response, bytes):
                    extracted_text = response.decode('utf-8', errors='replace')
                else:
                    extracted_text = str(response)
                
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                # Word文档 - 返回提示信息
                extracted_text = f"""
{file_name}

[Word文档内容预览]

此文档是Microsoft Word格式({mime_type})。
由于技术限制，我们无法在浏览器中显示完整的格式化内容。

您可以使用以下方式查看完整内容:
1. 在Google Drive中打开此文档
2. 下载此文档后在Microsoft Word中打开

文档基本信息:
- 名称: {file_name}
- 类型: {mime_type}
- ID: {file_id}
                """
                
            elif mime_type.startswith('text/'):
                # 纯文本文件
                request = drive_service.files().get_media(fileId=file_id)
                file_content = request.execute()
                
                if isinstance(file_content, bytes):
                    extracted_text = file_content.decode('utf-8', errors='replace')
                else:
                    extracted_text = str(file_content)
                    
            elif mime_type == 'application/pdf':
                # PDF文件提示信息
                extracted_text = f"""
{file_name}

[PDF文档内容]

此文档是PDF格式，您可以在主预览窗口中直接查看内容。
您可以使用鼠标选择并复制PDF中的文本内容。

要查看完整PDF文档:
1. 在Google Drive中打开
2. 下载此PDF文件
                """
            
            else:
                # 其他不支持的类型
                extracted_text = f"""
{file_name}

无法提取此类型文件的文本内容: {mime_type}

请尝试在Google Drive中打开或下载此文件查看。
                """
                
            return jsonify({"text": extracted_text})
            
        except Exception as e:
            logger.error(f"提取文本出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": str(e)}), 500

    @app.route('/api/extract-pdf-text', methods=['POST'])
    def extract_pdf_text():
        """提取PDF文本内容"""
        try:
            data = request.json
            credentials = create_credentials_from_token(data)
            file_id = data.get('file_id')
            
            if not file_id:
                return jsonify({"error": "缺少file_id参数"}), 400
            
            # 创建临时文件
            temp_dir = tempfile.mkdtemp()
            pdf_path = os.path.join(temp_dir, f'temp_pdf_{file_id}.pdf')
            
            try:
                # 下载PDF文件
                drive_service = googleapiclient.discovery.build('drive', 'v3', credentials=credentials)
                request = drive_service.files().get_media(fileId=file_id)
                
                with open(pdf_path, 'wb') as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        status, done = downloader.next_chunk()
                
                # 使用PyPDF2提取文本
                text = ""
                with open(pdf_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num in range(len(pdf_reader.pages)):
                        page_text = pdf_reader.pages[page_num].extract_text() or ""
                        text += page_text + "\n\n"
                
                return jsonify({"text": text})
                
            finally:
                # 清理临时文件
                try:
                    os.remove(pdf_path)
                    os.rmdir(temp_dir)
                except:
                    pass
                    
        except Exception as e:
            logger.error(f"提取PDF文本时出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": str(e)}), 500

    @app.route('/api/pdf-preview-images', methods=['POST'])
    def get_pdf_preview_images():
        """生成PDF页面的预览图像"""
        try:
            data = request.json
            credentials = create_credentials_from_token(data)
            file_id = data.get('file_id')
            
            if not file_id:
                return jsonify({"error": "缺少file_id参数"}), 400
            
            # 创建临时文件
            temp_dir = tempfile.mkdtemp()
            pdf_path = os.path.join(temp_dir, f'temp_pdf_{file_id}.pdf')
            
            try:
                # 下载PDF文件并处理
                drive_service = googleapiclient.discovery.build('drive', 'v3', credentials=credentials)
                request = drive_service.files().get_media(fileId=file_id)
                
                with open(pdf_path, 'wb') as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        status, done = downloader.next_chunk()
                
                # 尝试生成预览图像
                preview_images = []
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=10)
                    
                    for img in images:
                        img_buffer = BytesIO()
                        img.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        base64_img = base64.b64encode(img_buffer.read()).decode('utf-8')
                        preview_images.append(base64_img)
                        
                except ImportError:
                    return jsonify({"error": "服务器未安装PDF预览所需库，请使用提取文本功能"}), 500
                
                return jsonify({"preview_images": preview_images})
                
            finally:
                # 清理临时文件
                try:
                    os.remove(pdf_path)
                    os.rmdir(temp_dir)
                except:
                    pass
                    
        except Exception as e:
            logger.error(f"生成PDF预览时出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": str(e)}), 500

    @app.route('/api/convert-to-html', methods=['POST'])
    def convert_to_html():
        """将文档转换为HTML"""
        try:
            data = request.json
            if not data:
                return jsonify({"error": "缺少请求数据"}), 400
                
            file_id = data.get('file_id')
            if not file_id:
                return jsonify({"error": "缺少file_id参数"}), 400
            
            # 创建凭证和服务
            credentials = create_credentials_from_token(data)
            drive_service = googleapiclient.discovery.build('drive', 'v3', credentials=credentials)
            
            # 获取文件元数据
            file_metadata = drive_service.files().get(fileId=file_id, fields="name,mimeType").execute()
            mime_type = file_metadata.get('mimeType', '')
            file_name = file_metadata.get('name', 'unknown')
            
            # 根据文档类型处理
            html_content = ""
            
            if mime_type == 'application/vnd.google-apps.document':
                # Google文档 - 直接导出为HTML
                response = drive_service.files().export(
                    fileId=file_id,
                    mimeType='text/html'
                ).execute()
                
                if isinstance(response, bytes):
                    html_content = response.decode('utf-8', errors='replace')
                else:
                    html_content = str(response)
                    
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                # Word文档 - 尝试转换
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                temp_file.close()
                
                try:
                    # 下载文件
                    request = drive_service.files().get_media(fileId=file_id)
                    with open(temp_file.name, 'wb') as f:
                        downloader = MediaIoBaseDownload(f, request)
                        done = False
                        while not done:
                            _, done = downloader.next_chunk()
                    
                    # 尝试使用python-docx转换为HTML
                    try:
                        import docx
                        doc = docx.Document(temp_file.name)
                        
                        # 简单HTML转换
                        html_parts = ['<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body>']
                        
                        # 添加标题
                        if file_name:
                            html_parts.append(f'<h1>{file_name}</h1>')
                        
                        # 转换段落
                        for para in doc.paragraphs:
                            if not para.text.strip():
                                html_parts.append('<p>&nbsp;</p>')
                            else:
                                html_parts.append(f'<p>{para.text}</p>')
                        
                        html_parts.append('</body></html>')
                        html_content = ''.join(html_parts)
                        
                    except ImportError:
                        # 如果python-docx不可用，返回基本HTML预览
                        html_content = f"""
                        <!DOCTYPE html>
                        <html>
                        <head>
                            <meta charset="UTF-8">
                            <style>
                                body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                                .note {{ background: #fffde7; padding: 15px; border-left: 4px solid #ffd600; margin: 20px 0; }}
                            </style>
                        </head>
                        <body>
                            <h1>{file_name}</h1>
                            <div class="note">
                                <p>此文档预览使用简化格式显示。原始文档的格式、图像和某些内容可能未显示。</p>
                                <p>请使用"在Drive中查看"或"下载文档"选项查看完整格式。</p>
                            </div>
                            <p>文档类型: {mime_type}</p>
                            <p>由于无法进行完整转换，仅显示基本信息。</p>
                        </body>
                        </html>
                        """
                        
                finally:
                    # 清理临时文件
                    try:
                        os.unlink(temp_file.name)
                    except:
                        pass
                        
            else:
                # 其他不支持的类型
                return jsonify({"error": f"不支持转换此类型: {mime_type}"}), 400
                
            # 清理HTML内容 (移除脚本和可能的危险内容)
            html_content = sanitize_html(html_content)
                
            return jsonify({"html": html_content})
            
        except Exception as e:
            logger.error(f"转换文档出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": str(e)}), 500

    @app.route('/api/export-document', methods=['POST'])
    def export_document():
        """导出Google文档为HTML格式"""
        try:
            token_info = request.json
            file_id = request.json.get('file_id')
            
            if not file_id:
                return jsonify({"error": "缺少file_id参数"}), 400
            
            credentials = create_credentials_from_token(token_info)
            drive_service = googleapiclient.discovery.build('drive', 'v3', credentials=credentials)
            
            # 导出文档
            response = drive_service.files().export(
                fileId=file_id,
                mimeType='text/html'
            ).execute()
            
            logger.info(f"导出文档成功，文件ID: {file_id}")
            
            return jsonify({
                "content": response.decode('utf-8') if isinstance(response, bytes) else response
            })
        except Exception as e:
            logger.error(f"导出文档时出错: {str(e)}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/extract-document-text', methods=['POST'])
    def extract_document_text():
        """提取文档文本内容"""
        try:
            data = request.json
            credentials = create_credentials_from_token(data)
            file_id = data.get('file_id')
            mime_type = data.get('mime_type')
            
            if not file_id:
                return jsonify({"error": "缺少file_id参数"}), 400
            
            drive_service = googleapiclient.discovery.build('drive', 'v3', credentials=credentials)
            
            # 根据MIME类型处理不同的文档
            if mime_type == 'application/vnd.google-apps.document':
                # Google Docs - 导出为纯文本
                response = drive_service.files().export(
                    fileId=file_id,
                    mimeType='text/plain'
                ).execute()
                
                if isinstance(response, bytes):
                    text = response.decode('utf-8')
                else:
                    text = response
                    
            elif mime_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                # Word文档处理
                temp_dir = tempfile.mkdtemp()
                doc_path = os.path.join(temp_dir, f'temp_doc_{file_id}.docx')
                
                try:
                    # 下载文件
                    request = drive_service.files().get_media(fileId=file_id)
                    with open(doc_path, 'wb') as f:
                        downloader = MediaIoBaseDownload(f, request)
                        done = False
                        while not done:
                            status, done = downloader.next_chunk()
                    
                    # 使用python-docx提取文本
                    try:
                        import docx
                        doc = docx.Document(doc_path)
                        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                        text = "\n\n".join(paragraphs)
                    except ImportError:
                        text = "服务器未安装Word文档处理库。请下载文件查看内容。"
                finally:
                    # 清理临时文件
                    try:
                        os.remove(doc_path)
                        os.rmdir(temp_dir)
                    except:
                        pass
            else:
                # 其他类型文档
                return jsonify({"error": "不支持的文档类型"}), 400
                
            return jsonify({"text": text})
            
        except Exception as e:
            logger.error(f"提取文档文本时出错: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": str(e)}), 500